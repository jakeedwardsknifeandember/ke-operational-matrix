import streamlit as st
import gspread
import json
import os
import google.generativeai as genai
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- SESSION STATE INITIALIZATION ---
if 'logged_in' not in st.session_state:
    st.session_state.logged_in = False

# --- COMPLIANCE DATABASE & AI INITIALIZATION ---
credentials = json.loads(st.secrets["gcp_service_account"])
gc = gspread.service_account_from_dict(credentials)

# Securely boots your API connection from your Streamlit vault
genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
model = genai.GenerativeModel('gemini-2.5-flash')

# Helper function to inject clean, uniform tracking borders around corporate header tables
def set_cell_border(cell, color="CCCCCC", sz="4", val="single"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), val)
        border.set(qn('w:sz'), sz)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

# ==========================================
# ADMINISTRATIVE GATEKEEPER
# ==========================================
if not st.session_state.logged_in:
    st.title("🪵 Knife & Ember Workspace")
    st.subheader("🔒 FSCO Administrative Authentication")
    
    password_input = st.text_input("System Password", type="password")
    if st.button("Authenticate Panel", use_container_width=True):
        if password_input == st.secrets["APP_PASSWORD"]:
            st.session_state.logged_in = True
            st.rerun()
        else:
            st.error("Access Denied: Invalid Administrative Token.")

# ==========================================
# ENTERPRISE AI ONBOARDING FACTORY
# ==========================================
if st.session_state.logged_in:
    st.title("🏭 FSMS Client Onboarding & AI Factory")
    st.markdown("Compile smart, complete compliance manuals matching your custom consulting design templates exactly.")
    st.divider()
    
    # --- STEP 1: ESTABLISHMENT BACKGROUND & CORE ARCHITECTURE ---
    st.header("🏢 1. Establishment Background & Core Architecture")
    col1, col2 = st.columns(2)
    with col1:
        client_name = st.text_input("Brand Name / Legal Entity Name:", placeholder="")
        client_location = st.text_input("Operational Unit Address / Province:", placeholder="")
    with col2:
        facility_type = st.selectbox(
            "Facility Operational Classification:",
            [
                "Central Commissary Kitchen", 
                "Full-Service Dine-In Restaurant", 
                "Cafe / Specialty Coffee Shop (Light Cooking)", 
                "Micro-Kiosk / Street-Side Retail Stand", 
                "Pastry Shop & Bakery (No Raw Protein Prep)"
            ]
        )
        regulatory_scope = st.selectbox(
            "Primary Regulatory Oversight Framework:",
            [
                "Local Government LGU Sanitation Code (Standard Retail)", 
                "FDA GHP/HACCP Mandatory Scope (Manufacturing/Commissaries)", 
                "NMIS Meat Inspection Enforcement (Primary Meat Processing)"
            ]
        )

    # --- STEP 2: INFRASTRUCTURE & STRUCTURAL UTILITIES ---
    st.header("🏗️ 2. Infrastructure & Structural Utilities")
    
    with st.expander("🛠️ Countertop Fabrication & Cold-Chain Shelving Materials", expanded=True):
        st.markdown("**Worktop & Prep Surface Fabrication:**")
        c_ss = st.checkbox("Food-Grade 304 Stainless Steel (Standard Corrosion Resistance)")
        c_stone = st.checkbox("Natural Stone / Marble & Granite (Porous, Temp-Retaining Dough Prep)")
        c_wood = st.checkbox("Hardwood / Butcher Block Prep Counters (Highly Porous Artisarial Surfaces)")
        
        st.markdown("---")
        st.markdown("**Cold-Chain Storage Shelving Units:**")
        s_epoxy = st.checkbox("Epoxy-Coated / Plastic Composite Shelving (High-Moisture Coolers)")
        s_chrome = st.checkbox("Chrome-Plated Wire Shelving (Dry Storage Only)")

    with st.expander("🚰 Environmental & Water Engineering Toggles", expanded=False):
        u_open = st.checkbox("Open-Air / Street-Facing Kiosk Facility Environment (Exposed Vector Risks)")
        u_manual_water = st.checkbox("Containerized / Manual Gravity Water Supply (No Connected Mains Plumbing)")
        u_ice = st.checkbox("On-Site Commercial Ice Production Equipment (High Slime/Mold Biofilm Risks)")
        u_hoods = st.checkbox("Commercial Ventilation Hoods & Active ANSUL Fire Suppression Arrays")

    # --- STEP 3: DRINK CAPABILITIES & BEVERAGE MATRIX ---
    st.header("🥤 3. Beverage Capabilities Matrix")
    with st.expander("☕ Specialty Beverage Program Toggles", expanded=False):
        d_coffee = st.checkbox("Coffee Based Operations (Espresso Lines / Steam-Wand Management)")
        d_tea = st.checkbox("Tea Based Programs (Bulk Brewed Batching / Infused Syrups)")
        d_milk = st.checkbox("Milk Based / Fresh Liquid Dairy Creamers (High-Risk Temperature Loops)")
        d_nondairy = st.checkbox("Non-Dairy Alternative Milks (Almond, Soy, Oat Allergen Segregation)")
        d_frappe = st.checkbox("Frappes / Milkshakes served with Whipping Cream Siphon Canisters")
        d_soda = st.checkbox("Soda Based Operations (Pressurized CO2 Gas Lines / Post-Mix Systems)")

    # --- STEP 4: CULINARY MENU CATEGORIES MATRIX ---
    st.header("🍕 4. Culinary Menu Categories Matrix")
    with st.expander("🍔 Active Food Menu Item Classifications", expanded=True):
        f_app = st.checkbox("Appetizers & Finger Foods (Nachos, Fries, Quesadillas)")
        f_rice = st.checkbox("Rice Bowls / Mass Grains Batching (Bacillus Cereus Mitigation Scope)")
        f_pasta = st.checkbox("Pre-boiled Pasta Handling & Starch Lines")
        f_pizza = st.checkbox("Pizza Production (High Flour Dust Allergen & Deck Oven Risks)")
        f_pastry = st.checkbox("Pastries & Pre-Baked Confectionery Goods")
        f_icecream = st.checkbox("Soft Serve Ice Cream Operations (High-Risk Liquid Hopper Wash Loops)")
        f_sandwich = st.checkbox("Sandwiches & Cold Ready-To-Eat (RTE) Manual Assemblies")
        f_burger = st.checkbox("Burgers & Flat-Top Griddle Operations (High-Velocity Ground Proteins)")
        f_salad = st.checkbox("Salads & Fresh Raw Produce Washes (Surface Pathogen Controls)")
        f_poultry = st.checkbox("Chicken Processing (Raw Poultry Cross-Contamination Boundaries)")
        f_meat = st.checkbox("Pork or Beef Whole Muscle Cuts")
        f_seafood = st.checkbox("Seafood & Raw/Chilled Marine Proteins")

    st.divider()

    # --- STEP 5: AUTOMATION EXECUTION ---
    st.header("🚀 5. Execute AI Document Generation")
    if st.button("🔥 Run AI Compilation Engine", use_container_width=True):
        if not client_name.strip() or not client_location.strip():
            st.error("Compilation Halted: Brand Name and Address parameters cannot be empty.")
        else:
            formatted_sheet_name = client_name.strip().replace(" ", "_")
            master_path = "templates/master_core_fsms.docx"
            
            if not os.path.exists(master_path):
                st.error("Compilation Stopped: master_core_fsms.docx was not located inside your templates directory.")
            else:
                # --- PHASE A: READ MASTER DATA FOR AI CONTEXT ---
                with st.spinner("Analyzing master_core_fsms.docx structure..."):
                    try:
                        template_doc = Document(master_path)
                        core_text = "\n".join([p.text for p in template_doc.paragraphs if p.text.strip()])
                    except Exception as e:
                        st.error(f"Failed to read master layout text: {e}")
                        st.stop()

                # --- PHASE B: ASSEMBLE MODULAR COMPILATION QUEUE ---
                # Foundational core targets generated for all food businesses
                modules_to_generate = [
                    "PRP-01: Personnel Hygiene, Dermal Wound Control, and Clinical Illness Reporting",
                    "PRP-02: Facility Chemical Sanitation, Mechanical Warewashing, and Surface Efficacy Testing",
                    "PRP-03: Integrated Pest Management, Structural Exclusion Barriers, and Vector Data Monitoring",
                    "PRP-04: Environmental Waste Systems, Grease Interceptor Maintenance, and Wastewater Management",
                    "SOP-01: Inbound Logistics Receiving, Cold-Chain Validation, and Rejected Shipment Protocols",
                    "SOP-02: Inventory Cold Storage Systems, Cross-Contamination Hierarchies, and Allergen Segregation"
                ]

                # Dynamically append specialized operational modules based on checkbox arrays
                if u_ice: modules_to_generate.append("SOP-03: Commercial Ice Machine Mechanical Breakdown, Descaling, and Mold Decontamination")
                if d_coffee or d_milk or d_frappe: modules_to_generate.append("SOP-04: High-Risk Liquid Dairy Handling, Steam Wand Sanitation, and Espresso Line Maintenance")
                if d_nondairy: modules_to_generate.append("SOP-05: Plant-Based Milk Management and Cross-Contact Allergen Isolation")
                if f_rice or f_pasta: modules_to_generate.append("SOP-06: Cooked Starch and Mass Grain Stabilization (Bacillus Cereus Spore Prevention)")
                if f_pizza or f_pastry: modules_to_generate.append("SOP-07: Industrial Bakery Operations, Flour Dust Management, and Wheat Allergen Control")
                if f_icecream: modules_to_generate.append("SOP-08: Soft Serve Ice Cream Production, Hopper Hygiene, and Daily Teardown Sanatization")
                if f_sandwich: modules_to_generate.append("SOP-09: Ready-To-Eat Cold Line Assembly and Knife/Glove Manipulation Frequencies")
                if f_burger: modules_to_generate.append("SOP-10: Flat-Top Griddle Operations, Ground Protein Thermal Tracking, and Carbon Residue Scraping")
                if f_salad: modules_to_generate.append("SOP-11: Raw Produce Treatment, Chemical Antimicrobial Chlorination Wash, and Titration Controls")
                if f_poultry: modules_to_generate.append("SOP-12: Raw Poultry De-boning, Breading Line Separation, and Salmonella Ingress Barriers")
                if f_seafood: modules_to_generate.append("SOP-13: Chilled Marine Proteins, Histamine Accumulation Limits, and Parasite Destruction Tracking")

                # --- PHASE C: RUN MODULAR ASSEMBLY MULTI-PROMPT LOOP ---
                final_compiled_text_blocks = []
                progress_bar = st.progress(0)
                total_mods = len(modules_to_generate)

                for idx, target_module in enumerate(modules_to_generate):
                    with st.spinner(f"Compiling detailed standalone manual file: {target_module}..."):
                        
                        loop_prompt = f"""
                        You are an expert Food Safety Compliance Officer operating under Philippine regulations (RA 10611, FDA, and NMIS guidelines).
                        Your task is to write a highly detailed, uncompressed compliance manual section for {target_module} tailored specifically for our client.

                        CLIENT DATA:
                        - Client Name: {client_name}
                        - Location: {client_location}
                        - Classification: {facility_type}
                        - Oversight Framework: {regulatory_scope}
                        - Checked Setup Conditions: Surface Material options matching ({c_ss=}, {c_stone=}, {c_wood=}), Water system ({u_manual_water=}), Open kiosk layout ({u_open=}).

                        STYLE BENCHMARK TEXT CONTEXT:
                        {core_text[:3000]}

                        ABSOLUTE FORMATTING COMMANDS:
                        1. Generate content ONLY for the single module specified: {target_module}. Do not summarize or touch other areas.
                        2. Write a comprehensive manual section from scratch. You must include the full numbered prefix for all 9 core structural headings exactly as follows:
                           1. Purpose
                           2. Scope
                           3. Definitions
                           4. Responsibility
                           5. Procedure
                           6. Monitoring
                           7. Corrective Action
                           8. Verification
                           9. Records
                        3. Format subheadings inside section 5 with clear subnumbers (e.g., '5.1 Specific Execution Rules', '5.2 Frequency Parameters').
                        4. For specific operational item text lines under procedures, use a plain keyword label prefix followed by a colon (e.g., 'Uniforms: All staff must...', 'Method: Staff must scrub...').
                        5. Incorporate precise, actionable operational parameters tailored to this client's profile. (For example, if handling burgers or griddles, specify thermal cooking parameters of 155°F/68°C; if handling chicken, specify 165°F/74°C; if manual chemical sanitization is tracked, note Quat titration requirements of 200–400 ppm or Chlorine ranges of 50–100 ppm).
                        6. DO NOT use markdown characters, asterisks (**), or hashtags (##). 
                        7. DO NOT append any source brackets, bibliographies, footnotes, or numbers like .
                        """

                        try:
                            response = model.generate_content(loop_prompt)
                            module_text = response.text.strip()
                            
                            # Add delimiters to guide the Word formatting engine layout stitcher
                            final_compiled_text_blocks.append("[PAGE_BREAK]")
                            final_compiled_text_blocks.append(target_module)
                            final_compiled_text_blocks.append(module_text)
                        except Exception as e:
                            st.error(f"Engine Loop Error on {target_module}: {e}")
                            st.stop()
                        
                        progress_bar.progress((idx + 1) / total_mods)

                # --- PHASE D: RECONSTRUCT DOCX LAYER & APPLY OVERRIDES ---
                with st.spinner("Executing run-level layout overrides and margin justification (Times New Roman 9pt)..."):
                    try:
                        final_doc = Document(master_path)
                        
                        # Wipe out old placeholder text blocks safely from the XML DOM tree
                        while len(final_doc.paragraphs) > 0:
                            p_to_del = final_doc.paragraphs[0]
                            p_to_del._element.getparent().remove(p_to_del._element)
                        
                        # Flatten the text array array list blocks into simple paragraph strings
                        flat_raw_lines = "\n".join(final_compiled_text_blocks).split("\n")
                        
                        # Injects the thin-bordered standalone metadata card header block at the top of each page
                        def inject_corporate_header(doc_obj, title_text):
                            tbl = doc_obj.add_table(rows=2, cols=2)
                            tbl.autofit = False
                            tbl.columns[0].width = Inches(4.5)
                            tbl.columns[1].width = Inches(2.0)
                            
                            cell_00 = tbl.cell(0, 0)
                            p_00 = cell_00.paragraphs[0]
                            p_00.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            r_00 = p_00.add_run(f"🏢 FOOD SAFETY MANAGEMENT SYSTEM | {client_name.upper()}")
                            r_00.font.name = 'Times New Roman'
                            r_00.font.size = Pt(8)
                            r_00.font.color.rgb = RGBColor(100, 100, 100)
                            
                            cell_01 = tbl.cell(0, 1)
                            p_01 = cell_01.paragraphs[0]
                            p_01.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            r_01 = p_01.add_run("Doc Ref: FSMS-COLLECTION-2026")
                            r_01.font.name = 'Times New Roman'
                            r_01.font.size = Pt(8)
                            r_01.font.color.rgb = RGBColor(100, 100, 100)
                            
                            cell_10 = tbl.cell(1, 0)
                            p_10 = cell_10.paragraphs[0]
                            p_10.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            r_10 = p_10.add_run(title_text.upper())
                            r_10.font.name = 'Times New Roman'
                            r_10.font.size = Pt(10)
                            r_10.bold = True
                            
                            cell_11 = tbl.cell(1, 1)
                            p_11 = cell_11.paragraphs[0]
                            p_11.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            r_11 = p_11.add_run("Standard: RA 10611 / GHP")
                            r_11.font.name = 'Times New Roman'
                            r_11.font.size = Pt(8)
                            r_11.font.italic = True
                            r_11.font.color.rgb = RGBColor(100, 100, 100)
                            
                            for row in tbl.rows:
                                for cell in row.cells:
                                    set_cell_border(cell)
                            
                            spacer = doc_obj.add_paragraph()
                            spacer.paragraph_format.space_after = Pt(12)

                        is_next_line_title = False
                        is_first_page = True
                        
                        for line in flat_raw_lines:
                            cleaned_line = line.strip()
                            if not cleaned_line:
                                continue
                            
                            # Intercept page split signals
                            if cleaned_line == "[PAGE_BREAK]":
                                if not is_first_page:
                                    final_doc.add_page_break()
                                is_first_page = False
                                is_next_line_title = True
                                continue
                            
                            # Anchor tracking header card definitions
                            if is_next_line_title:
                                inject_corporate_header(final_doc, cleaned_line)
                                is_next_line_title = False
                                continue
                            
                            # INTERCEPTOR: Strips away any stray markdown formatting characters entirely
                            cleaned_line = cleaned_line.replace("***", "").replace("**", "").replace("*", "").replace("##", "").replace("#", "")
                            
                            # Identify section titles and subheaders (e.g., "1. Purpose" or "5.1 ")
                            is_heading = False
                            first_word = cleaned_line.split(" ")[0] if " " in cleaned_line else cleaned_line
                            if (first_word and first_word[0].isdigit() and "." in first_word) or \
                               cleaned_line.startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.")) or \
                               cleaned_line in ["Purpose", "Scope", "Definitions", "Responsibility", "Procedure", "Monitoring", "Corrective Action", "Verification", "Records"]:
                                is_heading = True
                            
                            new_p = final_doc.add_paragraph()
                            new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            
                            if is_heading:
                                run = new_p.add_run(cleaned_line)
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(9)
                                run.bold = True
                                new_p.paragraph_format.space_before = Pt(14)
                                new_p.paragraph_format.space_after = Pt(4)
                            
                            elif ":" in cleaned_line and not cleaned_line.startswith("http"):
                                # Inline Parameter Layout Engine (e.g., "Uniforms: All staff...")
                                label_part, text_part = cleaned_line.split(":", 1)
                                
                                label_run = new_p.add_run(label_part + ":")
                                label_run.font.name = 'Times New Roman'
                                label_run.font.size = Pt(9)
                                label_run.bold = True
                                
                                text_run = new_p.add_run(text_part)
                                text_run.font.name = 'Times New Roman'
                                text_run.font.size = Pt(9)
                                text_run.bold = False
                                
                                # Adds clean breathing gaps before inline fields so they never run together compressed
                                new_p.paragraph_format.space_before = Pt(8)
                                new_p.paragraph_format.space_after = Pt(3)
                                
                            else:
                                # Standard procedural sentences
                                run = new_p.add_run(cleaned_line)
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(9)
                                run.bold = False
                                new_p.paragraph_format.space_before = Pt(0)
                                new_p.paragraph_format.space_after = Pt(4)

                        # Save output and expose secure download button portal
                        output_filename = f"{client_name.strip().replace(' ', '_')}_Custom_FSMS.docx"
                        final_doc.save(output_filename)
                        st.success(f"🟢 Full-scope customized FSMS manual compiled for {client_name}!")
                        
                        with open(output_filename, "rb") as file:
                            st.download_button(
                                label=f"📥 Download Standalone Collection FSMS Manual for {client_name} (.docx)",
                                data=file,
                                file_name=output_filename,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
                        os.remove(output_filename)
                        
                    except Exception as e:
                        st.error(f"Layout Assembly Processing Failure: {e}")
